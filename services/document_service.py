import os
import platform
import shutil
import subprocess
import tempfile
from datetime import date, datetime
from pathlib import Path

from docx import Document


BULAN_INDONESIA = {
    1: "Januari",
    2: "Februari",
    3: "Maret",
    4: "April",
    5: "Mei",
    6: "Juni",
    7: "Juli",
    8: "Agustus",
    9: "September",
    10: "Oktober",
    11: "November",
    12: "Desember",
}


def format_tanggal_indonesia(value):
    if value in (None, ""):
        return ""

    if isinstance(value, str):
        text = value.strip()
        if not text:
            return ""

        parsed = None
        for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%d-%m-%Y", "%d/%m/%Y"):
            try:
                parsed = datetime.strptime(text, fmt)
                break
            except ValueError:
                pass

        if parsed is None:
            return text
        value = parsed

    if isinstance(value, (date, datetime)):
        return f"{value.day} {BULAN_INDONESIA[value.month]} {value.year}"

    return str(value)


def _replace_in_paragraph(paragraph, replacements):
    if not paragraph.runs:
        return

    original_text = "".join(run.text for run in paragraph.runs)
    if not original_text:
        return

    new_text = original_text
    for placeholder, value in replacements.items():
        if placeholder in new_text:
            new_text = new_text.replace(
                placeholder,
                "" if value is None else str(value),
            )

    if new_text == original_text:
        return

    # Taruh hasil pada run pertama agar format dasar paragraf tetap mengikuti template.
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _replace_in_table(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_in_paragraph(paragraph, replacements)
            for nested_table in cell.tables:
                _replace_in_table(nested_table, replacements)


def _replace_header_footer(document, replacements):
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            _replace_in_paragraph(paragraph, replacements)
        for table in section.header.tables:
            _replace_in_table(table, replacements)

        for paragraph in section.footer.paragraphs:
            _replace_in_paragraph(paragraph, replacements)
        for table in section.footer.tables:
            _replace_in_table(table, replacements)


def generate_document(template_path, output_path, replacements):
    template_path = os.path.abspath(template_path)
    output_path = os.path.abspath(output_path)

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template tidak ditemukan: {template_path}")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    document = Document(template_path)

    for paragraph in document.paragraphs:
        _replace_in_paragraph(paragraph, replacements)

    for table in document.tables:
        _replace_in_table(table, replacements)

    _replace_header_footer(document, replacements)
    document.save(output_path)

    if not os.path.exists(output_path):
        raise RuntimeError("Dokumen DOCX gagal dibuat.")

    return output_path


def _candidate_libreoffice_paths():
    candidates = []

    # 1. Bisa dipaksa lewat environment variable bila diperlukan.
    env_path = os.environ.get("LIBREOFFICE_PATH")
    if env_path:
        candidates.append(env_path)

    # 2. PATH sistem.
    for name in ("soffice", "soffice.exe", "libreoffice", "libreoffice.exe"):
        found = shutil.which(name)
        if found:
            candidates.append(found)

    if platform.system().lower() == "windows":
        env_dirs = [
            os.environ.get("ProgramFiles"),
            os.environ.get("ProgramFiles(x86)"),
            os.environ.get("LOCALAPPDATA"),
        ]

        for base in env_dirs:
            if not base:
                continue
            candidates.extend([
                os.path.join(base, "LibreOffice", "program", "soffice.exe"),
                os.path.join(base, "Programs", "LibreOffice", "program", "soffice.exe"),
            ])

        # Lokasi standar paling umum.
        candidates.extend([
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ])

        # Cari instalasi LibreOffice dengan versi/nama folder berbeda.
        for parent in (
            Path(r"C:\Program Files"),
            Path(r"C:\Program Files (x86)"),
        ):
            if parent.exists():
                for path in parent.glob("LibreOffice*/program/soffice.exe"):
                    candidates.append(str(path))

        # App Paths registry Windows.
        try:
            import winreg

            registry_locations = [
                (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Microsoft\Windows\CurrentVersion\App Paths\soffice.exe"),
                (winreg.HKEY_CURRENT_USER, r"SOFTWARE\Microsoft\Windows\CurrentVersion\App Paths\soffice.exe"),
                (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\WOW6432Node\Microsoft\Windows\CurrentVersion\App Paths\soffice.exe"),
            ]

            for hive, key_name in registry_locations:
                try:
                    with winreg.OpenKey(hive, key_name) as key:
                        value, _ = winreg.QueryValueEx(key, None)
                        if value:
                            candidates.append(value)
                except OSError:
                    pass
        except Exception:
            pass

    else:
        candidates.extend([
            "/usr/bin/libreoffice",
            "/usr/bin/soffice",
            "/usr/local/bin/libreoffice",
            "/usr/local/bin/soffice",
        ])

    # Hilangkan duplikat sambil mempertahankan urutan.
    result = []
    seen = set()
    for candidate in candidates:
        if not candidate:
            continue
        candidate = os.path.abspath(os.path.expandvars(os.path.expanduser(str(candidate))))
        key = candidate.lower() if platform.system().lower() == "windows" else candidate
        if key not in seen:
            seen.add(key)
            result.append(candidate)

    return result


def cari_libreoffice():

    # ========================================================
    # WINDOWS
    # ========================================================

    if os.name == "nt":

        path_windows = (
            r"C:\Program Files\LibreOffice"
            r"\program\soffice.exe"
        )

        print(
            "CEK LIBREOFFICE:",
            path_windows
        )

        print(
            "EXISTS:",
            os.path.exists(
                path_windows
            )
        )

        if os.path.isfile(
            path_windows
        ):
            return path_windows


    # ========================================================
    # COBA DARI PATH
    # ========================================================

    for command in [
        "soffice",
        "soffice.exe",
        "libreoffice",
    ]:

        ditemukan = shutil.which(
            command
        )

        if ditemukan:

            return ditemukan


    # ========================================================
    # LINUX / STREAMLIT CLOUD
    # ========================================================

    kandidat_linux = [
        "/usr/bin/libreoffice",
        "/usr/bin/soffice",
    ]

    for path in kandidat_linux:

        if os.path.isfile(
            path
        ):
            return path


    raise FileNotFoundError(
        "LibreOffice tidak ditemukan.\n"
        "Path yang diperiksa:\n"
        r"C:\Program Files\LibreOffice\program\soffice.exe"
    )


def convert_docx_to_pdf(docx_path, pdf_path):
    docx_path = os.path.abspath(docx_path)
    pdf_path = os.path.abspath(pdf_path)

    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"File DOCX tidak ditemukan: {docx_path}")

    soffice = cari_libreoffice()

    output_dir = os.path.dirname(pdf_path) or os.getcwd()
    os.makedirs(output_dir, exist_ok=True)

    # LibreOffice membuat PDF dengan nama dasar yang sama dengan DOCX.
    generated_pdf = os.path.join(
        output_dir,
        os.path.splitext(os.path.basename(docx_path))[0] + ".pdf",
    )

    # Hapus hasil lama agar kita yakin hasil berikut benar-benar baru.
    for old_path in {generated_pdf, pdf_path}:
        if os.path.exists(old_path):
            try:
                os.remove(old_path)
            except OSError:
                pass

    with tempfile.TemporaryDirectory(prefix="spd_lo_profile_") as profile_dir:
        profile_uri = Path(profile_dir).resolve().as_uri()

        command = [
            soffice,
            "--headless",
            "--nologo",
            "--nodefault",
            "--nolockcheck",
            "--nofirststartwizard",
            f"-env:UserInstallation={profile_uri}",
            "--convert-to",
            "pdf:writer_pdf_Export",
            "--outdir",
            output_dir,
            docx_path,
        ]

        result = subprocess.run(
            command,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=120,
            creationflags=(
                subprocess.CREATE_NO_WINDOW
                if platform.system().lower() == "windows"
                else 0
            ),
        )

    if result.returncode != 0 and not os.path.exists(generated_pdf):
        raise RuntimeError(
            "LibreOffice gagal mengonversi DOCX ke PDF.\n\n"
            f"Executable: {soffice}\n"
            f"Return code: {result.returncode}\n\n"
            f"STDOUT:\n{result.stdout}\n\n"
            f"STDERR:\n{result.stderr}"
        )

    if not os.path.exists(generated_pdf):
        raise RuntimeError(
            "LibreOffice selesai dijalankan tetapi file PDF tidak ditemukan.\n\n"
            f"Executable: {soffice}\n"
            f"DOCX: {docx_path}\n"
            f"PDF yang diharapkan: {generated_pdf}\n\n"
            f"STDOUT:\n{result.stdout}\n\n"
            f"STDERR:\n{result.stderr}"
        )

    if os.path.abspath(generated_pdf) != os.path.abspath(pdf_path):
        if os.path.exists(pdf_path):
            os.remove(pdf_path)
        shutil.move(generated_pdf, pdf_path)

    return pdf_path
